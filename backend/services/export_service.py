"""
Export service - Word, Excel, PDF report generation.
"""
import io
import logging
from datetime import datetime
from typing import Dict

logger = logging.getLogger(__name__)

RISK_COLORS = {
    'baixo': '70AD47',
    'moderado': 'FFD700',
    'alto': 'FF6600',
    'critico': 'FF0000',
    'sem_dados': 'CCCCCC',
}


class ExportService:
    """Generates reports in various formats."""

    def gerar_excel_dashboard(self, campaign, filters: dict = None) -> bytes:
        """Generate Excel dashboard report."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter

        from selectors.dashboard_selectors import DashboardSelectors
        from services.score_service import TIPO_DIMENSAO
        from services.risk_service import RiskService
        from services.score_service import score_service

        wb = Workbook()
        ws = wb.active
        ws.title = 'Dashboard'

        # Header
        ws['A1'] = f'VIVAMENTE 360º - Relatório de Resultados'
        ws['A2'] = f'Campanha: {campaign.nome}'
        ws['A3'] = f'Empresa: {campaign.empresa.nome}'
        ws['A4'] = f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
        ws['A1'].font = Font(bold=True, size=14)

        # Metrics
        metrics = DashboardSelectors.get_metrics(campaign, filters)
        ws['A6'] = 'INDICADORES GERAIS'
        ws['A6'].font = Font(bold=True)
        ws['A7'] = 'Total Convidados'
        ws['B7'] = metrics.get('total_convidados', 0)
        ws['A8'] = 'Total Respondentes'
        ws['B8'] = metrics.get('total_respondidos', 0)
        ws['A9'] = 'Taxa de Adesão (%)'
        ws['B9'] = metrics.get('taxa_adesao', 0)

        # Scores per dimension
        scores = DashboardSelectors.get_dimensoes_scores(campaign, filters)
        ws['A11'] = 'SCORES POR DIMENSÃO HSE-IT'
        ws['A11'].font = Font(bold=True)
        ws['A12'] = 'Dimensão'
        ws['B12'] = 'Score (0-4)'
        ws['C12'] = 'Tipo'
        ws['D12'] = 'Nível de Risco'

        row = 13
        for dim, score in scores.items():
            if score is None:
                continue
            tipo = TIPO_DIMENSAO.get(dim, 'positivo')
            nivel = score_service.classificar_risco(score, tipo)
            ws.cell(row=row, column=1).value = dim
            ws.cell(row=row, column=2).value = score
            ws.cell(row=row, column=3).value = tipo
            ws.cell(row=row, column=4).value = nivel
            # Color-code risk
            fill_color = RISK_COLORS.get(nivel, 'CCCCCC')
            ws.cell(row=row, column=4).fill = PatternFill(
                start_color=fill_color, end_color=fill_color, fill_type='solid'
            )
            row += 1

        # Auto-size columns
        for col in range(1, 5):
            ws.column_dimensions[get_column_letter(col)].width = 25

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    def gerar_excel_risk_matrix(self, campaign) -> bytes:
        """Generate risk matrix Excel report."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter
        from services.risk_calculation_service import RiskCalculationService

        risk_calc = RiskCalculationService()
        matriz = risk_calc.gerar_matriz_completa(campaign, campaign.empresa.cnae)

        wb = Workbook()
        ws = wb.active
        ws.title = 'Matriz de Risco'

        ws['A1'] = 'VIVAMENTE 360º - Matriz de Riscos Psicossociais (NR-1)'
        ws['A2'] = f'Campanha: {campaign.nome}'
        ws['A3'] = f'Empresa: {campaign.empresa.nome}'
        ws['A4'] = f'CNAE: {campaign.empresa.cnae}'
        ws['A5'] = f'Total de Respondentes: {matriz.get("total_respostas", 0)}'
        ws['A1'].font = Font(bold=True, size=13)

        headers = ['Dimensão', 'Score HSE-IT', 'Probabilidade (P)', 'Severidade (S)', 'NR (P×S)', 'Nível de Risco']
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=7, column=col, value=h)
            cell.font = Font(bold=True)

        row = 8
        for fator in matriz.get('fatores', []):
            ws.cell(row=row, column=1).value = fator.get('dimensao', '')
            ws.cell(row=row, column=2).value = fator.get('score_hse', '')
            ws.cell(row=row, column=3).value = fator.get('probabilidade', '')
            ws.cell(row=row, column=4).value = fator.get('severidade', '')
            ws.cell(row=row, column=5).value = fator.get('nr_score', '')
            nivel = fator.get('nivel', 'sem_dados')
            ws.cell(row=row, column=6).value = nivel
            fill_color = RISK_COLORS.get(nivel, 'CCCCCC')
            ws.cell(row=row, column=6).fill = PatternFill(
                start_color=fill_color, end_color=fill_color, fill_type='solid'
            )
            row += 1

        for col in range(1, 7):
            ws.column_dimensions[get_column_letter(col)].width = 20

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    def gerar_word_relatorio(self, campaign) -> bytes:
        """Generate Word report."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from selectors.dashboard_selectors import DashboardSelectors
        from services.score_service import TIPO_DIMENSAO
        from services.risk_service import RiskService
        from services.score_service import score_service

        doc = Document()
        doc.add_heading('VIVAMENTE 360º - Relatório de Avaliação Psicossocial', 0)
        doc.add_paragraph(f'Campanha: {campaign.nome}')
        doc.add_paragraph(f'Empresa: {campaign.empresa.nome}')
        doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
        doc.add_paragraph('')

        metrics = DashboardSelectors.get_metrics(campaign)
        doc.add_heading('1. Indicadores de Participação', level=1)
        doc.add_paragraph(f'Total de convidados: {metrics.get("total_convidados", 0)}')
        doc.add_paragraph(f'Total de respondentes: {metrics.get("total_respondidos", 0)}')
        doc.add_paragraph(f'Taxa de adesão: {metrics.get("taxa_adesao", 0)}%')

        scores = DashboardSelectors.get_dimensoes_scores(campaign)
        doc.add_heading('2. Resultados por Dimensão HSE-IT', level=1)
        for dim, score in scores.items():
            if score is None:
                continue
            tipo = TIPO_DIMENSAO.get(dim, 'positivo')
            nivel = score_service.classificar_risco(score, tipo)
            doc.add_paragraph(f'{dim}: Score {score:.2f} - Nível: {nivel.upper()}')

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


export_service = ExportService()
